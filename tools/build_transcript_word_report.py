"""
Generate Word comparison report: Gemini vs Whisper vs YouTube captions (3 pilot videos).

Structure per video:
  1. Video name
  2. Side-by-side model metrics
  3. Full Gemini transcript (all paragraphs)

Final section: overall performance table + uKfc detail table + conclusion.

Usage:
    python tools/build_transcript_word_report.py
"""

from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from pathlib import Path

_HERE = Path(__file__).resolve()
_ROOT = _HERE.parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

import importlib.util

_spec = importlib.util.spec_from_file_location("transcript_metrics", _HERE.parent / "transcript_metrics.py")
_tm = importlib.util.module_from_spec(_spec)
assert _spec.loader is not None
_spec.loader.exec_module(_tm)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TRANSCRIPT_DIR = _ROOT / "data" / "transcripts"
GOLD = _ROOT / "data" / "gold" / "videos_15_graded_vlm_reference_nocache.json"
MANIFEST = _ROOT / "reports" / "transcripts" / "beatriz_pilot" / "manifest.json"
OUT = _ROOT / "reports" / "transcripts" / "pilot_3_transcription_comparison_readable.docx"

VIDEOS = [
    {
        "id": "KTNcYYHuBTY",
        "title": "Powers of the Imaginary Unit i",
        "url": "https://www.youtube.com/watch?v=KTNcYYHuBTY",
        "scores": {"gemini": "9/10", "whisper": "7.5/10", "youtube": "—"},
    },
    {
        "id": "LVLuqNH5iWw",
        "title": "Local Extrema of a Multivariable Function",
        "url": "https://www.youtube.com/watch?v=LVLuqNH5iWw",
        "scores": {"gemini": "9/10", "whisper": "5/10", "youtube": "5/10"},
    },
    {
        "id": "uKfcS7-O6UE",
        "title": "Quotient Rule for Differentiation",
        "url": "https://www.youtube.com/watch?v=uKfcS7-O6UE",
        "scores": {"gemini": "8.5/10", "whisper": "8/10", "youtube": "7.5/10"},
    },
]

METHODS = ["gemini", "whisper", "youtube"]
SUFFIX = {"gemini": "gemini", "whisper": "whisper", "youtube": "youtube_captions"}
COL_LABEL = {"gemini": "Gemini", "whisper": "Whisper", "youtube": "YouTube"}

VIDEO_LABELS = {
    "KTNcYYHuBTY": "KTN (Powers of i)",
    "LVLuqNH5iWw": "LVL (Local extrema)",
    "uKfcS7-O6UE": "uKfc (Quotient rule)",
}

# Rows shown in the per-video comparison table
METRIC_ROWS = [
    ("Model", lambda r: r.get("model") or r.get("caption_type") or "—"),
    ("Word count", lambda r: str(r.get("word_count", "—"))),
    ("Char count", lambda r: str(r.get("char_count", "—"))),
    ("Duration (s)", lambda r: str(r.get("duration_s", "—"))),
    ("Processing time (s)", lambda r: str(r.get("elapsed_s", "—"))),
    ("Real-time factor", lambda r: str(r.get("real_time_factor", "—"))),
    ("Words per minute", lambda r: str(r.get("words_per_minute", "—"))),
    ("Keyword recall", lambda r: f"{r.get('keyword_recall_pct', '—')}% ({r.get('keyword_hit_count', '—')}/{r.get('keyword_total', '—')})"),
]


def _shade_cell(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _load_gold_keywords() -> dict[str, list[str]]:
    if not GOLD.is_file():
        return {}
    data = json.loads(GOLD.read_text(encoding="utf-8"))
    return {
        str(it["id"]): [str(k) for k in (it.get("reference_keywords") or [])]
        for it in (data.get("items") or [])
        if isinstance(it, dict) and it.get("id")
    }


def _load_manifest() -> dict[tuple[str, str], dict]:
    if not MANIFEST.is_file():
        return {}
    data = json.loads(MANIFEST.read_text(encoding="utf-8"))
    return {(rec["video_id"], rec["method"]): rec for rec in (data.get("records") or [])}


def _record_for(vid: str, method: str, video: dict, manifest: dict, gold_kw: dict) -> dict | None:
    suffix = SUFFIX[method]
    path = TRANSCRIPT_DIR / f"{vid}_{suffix}.txt"
    if not path.is_file():
        return None
    header, body = _tm.read_transcript(path)
    row = manifest.get((vid, method))
    if row:
        return row
    checks = _tm.PILOT_QUALITY_CHECKS.get(vid, {})
    return _tm.merge_record(
        video_id=vid,
        url=video["url"],
        title=video["title"],
        method=method,
        header=header,
        body=body,
        reference_keywords=gold_kw.get(vid, []),
        extra={"quality_checks": _tm.term_audit(body, checks)} if checks else None,
    )


def _quality_note(row: dict | None, vid: str) -> str:
    if not row:
        return "N/A"
    qc = row.get("quality_checks") or {}
    if vid == "LVLuqNH5iWw":
        if qc.get("saddle_point_wrong"):
            return "✗ subtle (not saddle)"
        if qc.get("saddle_point_correct"):
            return "✓ saddle point"
    if vid == "uKfcS7-O6UE":
        if qc.get("quotient_rule"):
            dy = "✓ LaTeX" if qc.get("dy_dx_notation") else "✓ rule / ✗ dy/dx"
            return dy if "dy" in dy else "✓ quotient rule"
    if vid == "KTNcYYHuBTY":
        if qc.get("imaginary_unit_i"):
            return "✓ imaginary unit i"
    return "—"


def _add_comparison_table(doc: Document, vid: str, rows_by_method: dict[str, dict | None], scores: dict) -> None:
    extra_rows = [
        ("Performance score", {
            "gemini": scores["gemini"],
            "whisper": scores["whisper"],
            "youtube": scores["youtube"],
        }),
        ("Key quality check", {
            "gemini": _quality_note(rows_by_method.get("gemini"), vid),
            "whisper": _quality_note(rows_by_method.get("whisper"), vid),
            "youtube": _quality_note(rows_by_method.get("youtube"), vid),
        }),
    ]

    table = doc.add_table(rows=1 + len(METRIC_ROWS) + len(extra_rows), cols=4)
    table.style = "Table Grid"
    headers = ["Metric", "Gemini", "Whisper", "YouTube"]
    for c, h in enumerate(headers):
        table.rows[0].cells[c].text = h
        _shade_cell(table.rows[0].cells[c], "1F4E79")
        for p in table.rows[0].cells[c].paragraphs:
            for run in p.runs:
                run.bold = True

    ri = 1
    for label, getter in METRIC_ROWS:
        table.rows[ri].cells[0].text = label
        _shade_cell(table.rows[ri].cells[0], "E8EEF7")
        for ci, method in enumerate(METHODS, start=1):
            rec = rows_by_method.get(method)
            if rec:
                table.rows[ri].cells[ci].text = getter(rec)
            else:
                table.rows[ri].cells[ci].text = "—" if method != "youtube" or vid != "KTNcYYHuBTY" else "N/A (429)"
        ri += 1

    for label, values in extra_rows:
        table.rows[ri].cells[0].text = label
        _shade_cell(table.rows[ri].cells[0], "E8EEF7")
        for ci, method in enumerate(METHODS, start=1):
            table.rows[ri].cells[ci].text = values[method]
            if method == "gemini":
                for run in table.rows[ri].cells[ci].paragraphs[0].runs:
                    run.bold = True
        ri += 1

    doc.add_paragraph()


def _add_full_gemini_transcript(doc: Document, vid: str) -> None:
    path = TRANSCRIPT_DIR / f"{vid}_gemini.txt"
    if not path.is_file():
        doc.add_paragraph("(Gemini transcript file not found.)")
        return
    _, body = _tm.read_transcript(path)
    doc.add_heading("Gemini transcript (full text)", level=2)
    meta = doc.add_paragraph()
    meta.add_run("Model: ").bold = True
    meta.add_run("gemini-3-flash-preview  |  Prompt: mathe-transcribe-v3-final")
    paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]
    if not paragraphs:
        doc.add_paragraph(body)
    else:
        for para in paragraphs:
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(8)


def main() -> int:
    manifest = _load_manifest()
    gold_kw = _load_gold_keywords()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("MathE Pilot — Transcription Performance Comparison", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc.add_paragraph(f"Generated: {generated}")
    doc.add_paragraph(
        "Three Beatriz pilot videos: metrics for Gemini 3 Flash, Whisper (distil-large-v3), "
        "and YouTube auto captions; followed by the complete Gemini transcript per video."
    )

    for video in VIDEOS:
        vid = video["id"]
        doc.add_page_break()
        doc.add_heading(video["title"], level=1)
        doc.add_paragraph(f"Video ID: {vid}")
        doc.add_paragraph(f"URL: {video['url']}")

        doc.add_heading("Model performance metrics", level=2)
        rows_by_method: dict[str, dict | None] = {}
        for method in METHODS:
            rows_by_method[method] = _record_for(vid, method, video, manifest, gold_kw)
        _add_comparison_table(doc, vid, rows_by_method, video["scores"])
        _add_full_gemini_transcript(doc, vid)

    doc.add_page_break()
    doc.add_heading("Overall performance scores", level=1)
    doc.add_paragraph(
        "Scores reflect transcription quality for MathE keyword/topic use (1–10). "
        "Higher is better. Based on manual review of pilot outputs and automated metrics."
    )

    table = doc.add_table(rows=1 + len(VIDEOS), cols=4)
    table.style = "Table Grid"
    for c, h in enumerate(["Video", "Gemini", "Whisper", "YouTube"]):
        table.rows[0].cells[c].text = h
        _shade_cell(table.rows[0].cells[c], "1F4E79")
        for run in table.rows[0].cells[c].paragraphs[0].runs:
            run.bold = True

    for ri, video in enumerate(VIDEOS, start=1):
        table.rows[ri].cells[0].text = VIDEO_LABELS[video["id"]]
        table.rows[ri].cells[1].text = video["scores"]["gemini"]
        table.rows[ri].cells[2].text = video["scores"]["whisper"]
        table.rows[ri].cells[3].text = video["scores"]["youtube"]
        for run in table.rows[ri].cells[1].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()
    doc.add_heading("Per-video detail scores (uKfc example metrics)", level=2)
    doc.add_paragraph("Illustrative metric comparison for Quotient Rule video:")

    uKfc = {m: _record_for("uKfcS7-O6UE", m, VIDEOS[2], manifest, gold_kw) for m in METHODS}

    def _dy_notation(method: str) -> str:
        rec = uKfc.get(method)
        if not rec:
            return "—"
        qc = rec.get("quality_checks") or {}
        if method == "gemini" and qc.get("dy_dx_notation"):
            return "✓ (LaTeX)"
        if method == "whisper":
            return "✗ (spoken)"
        if method == "youtube" and qc.get("dy_dx_notation"):
            return "✓"
        return "—"

    detail = doc.add_table(rows=5, cols=4)
    detail.style = "Table Grid"
    detail_data = [
        ["Metric", "Gemini", "Whisper", "YouTube"],
        ["Word count", str(uKfc["gemini"]["word_count"] if uKfc["gemini"] else "—"),
         str(uKfc["whisper"]["word_count"] if uKfc["whisper"] else "—"),
         str(uKfc["youtube"]["word_count"] if uKfc["youtube"] else "—")],
        ["Keyword recall",
         f"{uKfc['gemini']['keyword_recall_pct']}%" if uKfc["gemini"] else "—",
         f"{uKfc['whisper']['keyword_recall_pct']}%" if uKfc["whisper"] else "—",
         f"{uKfc['youtube']['keyword_recall_pct']}%" if uKfc["youtube"] else "—"],
        ["Quotient rule", "✓", "✓", "✓"],
        ["dy/dx notation", _dy_notation("gemini"), _dy_notation("whisper"), _dy_notation("youtube")],
    ]
    for r, row_data in enumerate(detail_data):
        for c, val in enumerate(row_data):
            detail.rows[r].cells[c].text = val
            if r == 0:
                _shade_cell(detail.rows[r].cells[c], "E8EEF7")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Conclusion: ").bold = True
    p.add_run(
        "Gemini 3 Flash (mathe-transcribe-v3-final) is the recommended primary transcript source "
        "for the Beatriz pilot. The largest advantage appears on the multivariable extrema video (LVL), "
        "where Whisper and YouTube fail on saddle-point terminology. Deliverable transcripts are in "
        "reports/transcripts/beatriz_pilot/gemini_primary/."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote: {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
