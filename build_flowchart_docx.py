"""One-off: build STEP_algorithm_flowchart.png + STEP_algorithm_flowchart.docx."""
from __future__ import annotations

import sys
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parent.parent
OUT_PNG = ROOT / "STEP_algorithm_flowchart.png"
OUT_DOCX = ROOT / "STEP_algorithm_flowchart.docx"


def _box(ax, xy, wh, text, face: str = "#eef2f7", edge: str = "#1e3a5f"):
    x, y = xy
    w, h = wh
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.03,rounding_size=0.12",
        facecolor=face,
        edgecolor=edge,
        linewidth=1.35,
        mutation_aspect=0.35,
    )
    ax.add_patch(patch)
    ax.text(
        x + w / 2,
        y + h / 2,
        text,
        ha="center",
        va="center",
        fontsize=8.2,
        color="#0f172a",
        wrap=True,
    )


def _arrow(ax, p0, p1):
    arr = FancyArrowPatch(
        p0,
        p1,
        arrowstyle="-|>",
        mutation_scale=12,
        linewidth=1.1,
        color="#334155",
        shrinkA=2,
        shrinkB=2,
    )
    ax.add_patch(arr)


def build_png() -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 14.2), dpi=150)
    ax.axis("off")

    w, x = 7.25, 1.35
    gap = 0.14
    cx = x + w / 2

    # (height, text, face) — stacked top → bottom (decreasing matplotlib y)
    rows: list[tuple[float, str, str]] = [
        (0.62, "Start: PDF, YouTube URL,\nor local video file", "#dbeafe"),
        (0.58, "Input routing (PDF vs video)", "#f1f5f9"),
        (0.62, "Cache lookup (SHA-256 / video id)", "#fef9c3"),
        (0.55, "Cache hit? → return stored result\nelse continue", "#fde68a"),
        (0.95, "PDF — Layer 0\nRasterize (adaptive DPI), raw text, page PNGs", "#e0e7ff"),
        (1.0, "PDF — Layers 2–3 (optional)\nNougat OCR · VLM LaTeX\n(normalize, retry, disk cache)", "#e0e7ff"),
        (0.82, "PDF — Layer 4 late fusion\nSingle solver prompt", "#e0e7ff"),
        (0.82, "PDF — Layer 5 LLM solve\nGemini / fallback · consensus", "#e0e7ff"),
        (0.78, "PDF — Layer 6 verify\nSymPy · final boxed answer line", "#e0e7ff"),
        (1.05, "PDF — Layer 1b + 7 tags\nRegex taxonomy scores · L7 pool\n+ free keywords (Gemini, T=0)", "#dcfce7"),
        (1.15, "Video — Layer 0v + 3v\nNormalize / upload · Quick: native VLM\nDeep: download + frame sampling", "#fce7f3"),
        (0.88, "Video Deep — group scenes\nTrigram Jaccard J = |∩|/|∪|\nmerge if J ≥ threshold", "#fce7f3"),
        (0.78, "Video Deep — batch keywords\nGemini: top-5 pool / scene (ordered)", "#fce7f3"),
        (0.62, "Video — taxonomy fallback\n(topic_from_keywords on pool overlap)", "#fbcfe8"),
        (0.62, "Output: JSON + web UI\nanswer · classification · keywords", "#bbf7d0"),
    ]

    y_top = 16.45
    positions: list[tuple[float, float]] = []
    y_cursor = y_top
    for h, *_ in rows:
        y_bottom = y_cursor - h
        positions.append((y_bottom, h))
        y_cursor = y_bottom - gap

    ax.set_xlim(0, 10)
    ax.set_ylim(y_cursor - 0.35, y_top + 0.35)

    for (y_b, h), (hi, txt, face) in zip(positions, rows):
        _box(ax, (x, y_b), (w, hi), txt, face)

    for i in range(len(positions) - 1):
        y0, h0 = positions[i]
        y1, h1 = positions[i + 1]
        _arrow(ax, (cx, y0), (cx, y1 + h1))

    fig.savefig(OUT_PNG, bbox_inches="tight", facecolor="white", edgecolor="none")
    plt.close(fig)
    return OUT_PNG


def build_docx() -> Path:
    doc = Document()

    t = doc.add_heading("STEP / MathE pipeline — flowchart and keyword ordering", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This note accompanies the system flowchart figure. It answers how keyword "
        "relevance and ordering are determined, and distinguishes rule-based scores "
        "from model-based ranking."
    )

    doc.add_heading("1. Keyword relevance and order (technical)", level=1)
    p = doc.add_paragraph()
    p.add_run("A. Closed pool and free-form tasks (Layer 7, ").italic = True
    p.add_run("keyword_eval.py").italic = True
    p.add_run("). ").italic = False
    doc.add_paragraph(
        "Task 1 asks the language model for five short English keywords; Task 2 "
        "restricts choices to the supplied pool. Both use temperature 0 and "
        "explicit instructions to list keywords from most to least relevant. "
        "The implementation does not compute a numeric relevance score in code: "
        "ordering comes from the model’s constrained generation. After the call, "
        "the reply is split on commas; no re-ranking or tie-break formula is applied.",
        style="List Bullet",
    )
    p = doc.add_paragraph()
    p.add_run("B. Taxonomy card (“Most relevant keywords”, ").italic = True
    p.add_run("taxonomy.py").italic = True
    p.add_run("). ").italic = False
    doc.add_paragraph(
        "Here each keyword rule has several regex patterns. For text x, the score "
        "hr(x) is the count of patterns in that rule that match x (case-insensitive). "
        "Keywords are sorted primarily by descending hr, secondarily by alphabetic "
        "name for ties. Subtopic selection uses the same counting idea at subtopic "
        "level: the subtopic with the largest number of matching subtopic patterns wins.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Deep video mode reuses the same pool ordering instructions in a batch prompt "
        "(layer3v_frames.py): again model-ordered, with post-processing only to enforce "
        "exact pool strings and deduplication.",
        style="List Bullet",
    )
    p = doc.add_paragraph()
    p.add_run("C. Frame grouping (Deep video only). ").bold = True
    doc.add_paragraph(
        "This step does not rank keywords. Normalized LaTeX strings are turned into "
        "character trigram sets T(s). Similarity between two strings is Jaccard index "
        "J(a,b) = |T(a) ∩ T(b)| / |T(a) ∪ T(b)|. Frames merge into one scene when J "
        "exceeds a fixed threshold (default 0.65).",
        style="List Bullet",
    )

    doc.add_heading("2. Flowchart", level=1)
    doc.add_paragraph(
        "The figure summarises the main stages from ingestion through tagging. "
        "PDF and video branches are stacked for readability; in deployment the "
        "web layer selects one branch per request."
    )
    doc.add_picture(str(OUT_PNG), width=Cm(16.5))

    doc.add_page_break()
    doc.add_heading("Akış şeması ve anahtar kelime önceliği (Türkçe)", level=0)

    doc.add_heading("Anahtar kelime alakası ve sıralama nasıl belirleniyor?", level=1)
    doc.add_paragraph(
        "Sistemde üç ayrı mekanizma vardır; bunların ikisi açık matematiksel veya "
        "kural tabanlı skor kullanır, biri ise dil modelinin talimatla ürettiği sırayı "
        "olduğu gibi kaydeder."
    )
    doc.add_paragraph(
        "Katman 7 (Task 1 ve Task 2): Kapalı havuzdan seçim ve serbest beş anahtar "
        "kelime tamamen Gemini çağrısı ile yapılır; sıcaklık 0’dır ve istemde "
        "“en alakalıdan en aza” sıralama açıkça istenir. Kod tarafında her anahtar "
        "kelime için ayrı bir skor denklemi hesaplanmaz; modelin çıktısı virgülle "
        "bölünür, ek bir yeniden sıralama yoktur.",
        style="List Number",
    )
    doc.add_paragraph(
        "Taxonomy (MathE tarzı kart): Her anahtar kelime kuralının birden çok "
        "düzenli ifadesi vardır; metinde eşleşen ifade sayısı hr ile özetlenir. "
        "Anahtar kelimeler önce hr büyükten küçüğe, eşitlikte ada göre sıralanır. "
        "Alt konu seçimi de benzer şekilde alt konu desenlerinin eşleşme sayısının "
        "en yüksek olduğu alt konuya verilir.",
        style="List Number",
    )
    doc.add_paragraph(
        "Video Derin mod: Kareleri birleştirmek için üçlü (trigram) Jaccard "
        "benzerliği kullanılır: J(a,b)=|T(a)∩T(b)|/|T(a)∪T(b)|. Bu, sahne "
        "oluşturma içindir; anahtar kelime önceliği yine toplu LLM çıktısının "
        "sırasına dayanır.",
        style="List Number",
    )

    doc.add_heading("Akış şeması", level=1)
    doc.add_paragraph(
        "Aşağıdaki şekil, örnek “flowchart examples.docx” belgesindeki gibi "
        "görsel bir blok diyagram olarak sunulmuştur (Word içinde PNG olarak gömülü)."
    )
    doc.add_picture(str(OUT_PNG), width=Cm(16.5))

    doc.save(OUT_DOCX)
    return OUT_DOCX


def main() -> int:
    build_png()
    build_docx()
    print("Wrote:", OUT_PNG)
    print("Wrote:", OUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
